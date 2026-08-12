import os
import shutil
from dotenv import load_dotenv
from docx import Document
from langchain_chroma import Chroma
from langchain_huggingface import HuggingFaceEmbeddings

load_dotenv(override=True)

from chunkers import BaselineChunker, StructureAwareChunker

def parse_recipes_from_docx(filepath):
    doc = Document(filepath)
    recipes = []
    
    # State tracking
    def reset_recipe():
        return {
            "source_file": os.path.basename(filepath),
            "recipe_id": "",
            "title": "",
            "description": "",
            "cuisine": "",
            "dietary_tags": "",
            "allergens": "",
            "yield": "",
            "ingredients_headers": [],
            "ingredients_rows": [],
            "method_steps": [],
            "notes": ""
        }
        
    current_recipe = reset_recipe()
    current_section = None
    
    # Helper to clean text
    def clean(t): return t.strip()

    # The docx structure starts with title page info, then recipes.
    # We identify a recipe start if we see a short paragraph without a prefix and no current_section active,
    # or by specifically looking for known recipe names from the dataset.
    # But let's build it dynamically based on the parsed structure.
    
    RECIPE_TITLES = ["Idli", "Dosa", "Ven Pongal", "Sambar", "Rasam", "Medu Vada"]

    for element in doc.element.body:
        if element.tag.endswith('p'):
            # paragraph
            for p in doc.paragraphs:
                if p._element == element:
                    text = clean(p.text)
                    if not text:
                        break
                    
                    if text == "TAMIL NADU RECIPE BOOK" or "Six classic" in text:
                        break # skip intro
                        
                    if text in RECIPE_TITLES:
                        if current_recipe.get("title"):
                            recipes.append(current_recipe)
                        current_recipe = reset_recipe()
                        current_section = None
                        current_recipe["title"] = text
                        current_recipe["recipe_id"] = text.lower().replace(" ", "_")
                        break
                    
                    if text.startswith("Cuisine:"):
                        current_recipe["cuisine"] = text.replace("Cuisine:", "").strip()
                    elif text.startswith("Dietary Tags:"):
                        current_recipe["dietary_tags"] = text.replace("Dietary Tags:", "").strip()
                    elif text.startswith("Allergens:"):
                        current_recipe["allergens"] = text.replace("Allergens:", "").strip()
                    elif text.startswith("Yield:"):
                        current_recipe["yield"] = text.replace("Yield:", "").strip()
                    elif text == "Ingredients":
                        current_section = "ingredients"
                    elif text == "Method":
                        current_section = "method"
                    elif text == "Cook's Notes":
                        current_section = "notes"
                    elif current_section == "method":
                        if text and text[0].isdigit():
                            current_recipe["method_steps"].append(text)
                        else:
                            current_recipe["method_steps"].append(text)
                    elif current_section == "notes":
                        current_recipe["notes"] += text + " "
                    else:
                        # Description
                        if not current_recipe.get("description") and current_recipe.get("title"):
                            current_recipe["description"] = text
                            
                    break
        elif element.tag.endswith('tbl'):
            # table
            for t in doc.tables:
                if t._element == element:
                    if current_section == "ingredients":
                        for r_idx, row in enumerate(t.rows):
                            row_data = [clean(cell.text) for cell in row.cells]
                            if r_idx == 0:
                                current_recipe["ingredients_headers"] = row_data
                            else:
                                current_recipe["ingredients_rows"].append(row_data)
                    break
                    
    # append the last one
    if current_recipe.get("title"):
        recipes.append(current_recipe)
        
    return recipes

def dict_to_text(recipe_data):
    """Convert recipe dict back to a flat string for the baseline chunker."""
    lines = []
    lines.append(recipe_data["title"])
    if recipe_data["description"]: lines.append(recipe_data["description"])
    if recipe_data["cuisine"]: lines.append(f"Cuisine: {recipe_data['cuisine']}")
    if recipe_data["dietary_tags"]: lines.append(f"Dietary Tags: {recipe_data['dietary_tags']}")
    if recipe_data["allergens"]: lines.append(f"Allergens: {recipe_data['allergens']}")
    if recipe_data["yield"]: lines.append(f"Yield: {recipe_data['yield']}")
    
    lines.append("Ingredients")
    if recipe_data["ingredients_headers"]:
        lines.append(" | ".join(recipe_data["ingredients_headers"]))
    for row in recipe_data["ingredients_rows"]:
        lines.append(" | ".join(row))
        
    lines.append("Method")
    for step in recipe_data["method_steps"]:
        lines.append(step)
        
    if recipe_data["notes"]:
        lines.append("Cook's Notes")
        lines.append(recipe_data["notes"])
        
    return "\n".join(lines)

def ingest():
    filepath = "TamilNadu_Recipe_Cards.docx"
    print(f"Reading {filepath}...")
    
    if not os.path.exists(filepath):
        print(f"File {filepath} not found!")
        return

    recipes = parse_recipes_from_docx(filepath)
    print(f"Parsed {len(recipes)} recipes.")
    
    # Init Chunkers
    baseline_chunker = BaselineChunker()
    struct_chunker = StructureAwareChunker()
    
    baseline_docs = []
    struct_docs = []
    
    for r in recipes:
        # Base metadata for baseline chunker
        meta = {
            "source_file": r["source_file"],
            "recipe_id": r["recipe_id"],
            "cuisine": r["cuisine"],
            "dietary_tags": r["dietary_tags"]
        }
        
        # 1. Baseline
        text_content = dict_to_text(r)
        b_docs = baseline_chunker.chunk(text_content, meta)
        baseline_docs.extend(b_docs)
        
        # 2. Structure-aware
        s_docs = struct_chunker.chunk(r)
        struct_docs.extend(s_docs)
        
    print(f"Generated {len(baseline_docs)} baseline chunks.")
    print(f"Generated {len(struct_docs)} structure-aware chunks.")
    
    # Store in Chroma
    embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
    
    # Clear existing DB if any to avoid duplicates
    db_path = "./chroma_db"
    if os.path.exists(db_path):
        print("Clearing existing Vector DB...")
        shutil.rmtree(db_path)
        
    print("Ingesting baseline docs into Chroma...")
    Chroma.from_documents(
        baseline_docs, 
        embeddings, 
        collection_name="recipes_baseline",
        persist_directory=db_path
    )
    
    print("Ingesting structure-aware docs into Chroma...")
    Chroma.from_documents(
        struct_docs, 
        embeddings, 
        collection_name="recipes_structure",
        persist_directory=db_path
    )
    
    print("Ingestion complete!")

if __name__ == "__main__":
    ingest()
