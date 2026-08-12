from docx import Document
import sys

def inspect_docx(filepath):
    print(f"Inspecting: {filepath}")
    doc = Document(filepath)
    
    for i, element in enumerate(doc.element.body):
        if element.tag.endswith('p'):
            # paragraph
            for p in doc.paragraphs:
                if p._element == element:
                    text = p.text.strip()
                    if text:
                        style_name = p.style.name if p.style else "None"
                        try:
                            # Use ascii encoding to avoid windows console print errors with emojis/special chars
                            print(f"[Paragraph] Style: {style_name} | Text: {text[:100].encode('ascii', 'ignore').decode('ascii')}")
                        except Exception as e:
                            print(f"[Paragraph] Style: {style_name} | Text: <unprintable characters>")
                    break
        elif element.tag.endswith('tbl'):
            # table
            for t in doc.tables:
                if t._element == element:
                    print(f"[Table] Rows: {len(t.rows)}, Cols: {len(t.columns) if t.rows else 0}")
                    for r_idx, row in enumerate(t.rows):
                        row_data = [cell.text.strip() for cell in row.cells]
                        print(f"  Row {r_idx}: {row_data}")
                    break

if __name__ == "__main__":
    filepath = sys.argv[1] if len(sys.argv) > 1 else "TamilNadu_Recipe_Cards.docx"
    inspect_docx(filepath)
